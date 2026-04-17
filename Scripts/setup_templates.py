import os
from docx import Document

def create_reporte_servicio():
    doc = Document()
    doc.add_heading('Reporte de Servicio', 0)
    
    doc.add_heading('Datos Generales', level=1)
    doc.add_paragraph('Empresa: {{ empresa }}')
    doc.add_paragraph('RFC: {{ rfc }} | Teléfono: {{ telefono }}')
    doc.add_paragraph('Ubicación: {{ municipio }}, {{ estado }}, CP: {{ cp }}')
    doc.add_paragraph('Fecha: {{ fecha }}')
    doc.add_paragraph('Técnico: {{ tecnico }}')
    
    doc.add_heading('Detalles del Servicio', level=1)
    doc.add_paragraph('El técnico {{ tecnico }} ha completado el servicio de manera satisfactoria.')
    
    doc.add_heading('Evidencia (Fachada)', level=1)
    doc.add_paragraph('{{ fachada }}')
    
    doc.save('plantillas/reporte_servicio.docx')

def create_contrato_min():
    doc = Document()
    doc.add_heading('Contrato de Servicio', 0)
    
    doc.add_heading('Datos Generales', level=1)
    doc.add_paragraph('Empresa: {{ empresa }}')
    doc.add_paragraph('RFC: {{ rfc }} | Teléfono: {{ telefono }}')
    doc.add_paragraph('Ubicación: {{ municipio }}, {{ estado }}, CP: {{ cp }}')
    doc.add_paragraph('Fecha de Inicio: {{ fecha }}')
    
    doc.add_heading('Condiciones Financieras', level=1)
    doc.add_paragraph('El monto acordado a pagar por los servicios prestados es de ${{ monto }} USD.')
    
    doc.add_heading('Anexo Fotográfico', level=1)
    doc.add_paragraph('{{ fachada }}')
    
    doc.save('plantillas/contrato_min.docx')

if __name__ == "__main__":
    if not os.path.exists('plantillas'):
        os.makedirs('plantillas')
    create_reporte_servicio()
    create_contrato_min()
    print("Plantillas generadas exitosamente en la carpeta 'plantillas/'.")
